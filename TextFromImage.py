import cv2
import numpy as np
from PIL import Image
import pytesseract
from docx import Document  # Import the Document class from python-docx

# Set up Tesseract path if needed (Windows)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Load the image
screenshot_path = r"D:\Pictures\Screenshots\Screenshot (77).png"  # Replace with your screenshot path
image = cv2.imread(screenshot_path)

# Variables to hold selected coordinates
start_point = None
end_point = None

def select_text(event, x, y, flags, param):
    global start_point, end_point

    if event == cv2.EVENT_LBUTTONDOWN:
        start_point = (x, y)
    
    elif event == cv2.EVENT_LBUTTONUP:
        end_point = (x, y)
        cv2.rectangle(image, start_point, end_point, (0, 255, 0), 2)
        cv2.imshow("Image", image)

# Create a window to display the image
cv2.namedWindow("Image")
cv2.setMouseCallback("Image", select_text)

# Display the image
cv2.imshow("Image", image)
cv2.waitKey(0)

# Ensure points are set
if start_point and end_point:
    x1, y1 = min(start_point[0], end_point[0]), min(start_point[1], end_point[1])
    x2, y2 = max(start_point[0], end_point[0]), max(start_point[1], end_point[1])
    
    if x1 < 0 or y1 < 0 or x2 > image.shape[1] or y2 > image.shape[0]:
        print("Selected region is out of image bounds.")
    else:
        selected_region = image[y1:y2, x1:x2]

        if selected_region.size == 0:
            print("Selected region is empty.")
        else:
            # Perform OCR on the selected region
            pil_image = Image.fromarray(cv2.cvtColor(selected_region, cv2.COLOR_BGR2RGB))
            extracted_text = pytesseract.image_to_string(pil_image)
            print(f"Extracted Text: {extracted_text.strip()}")

            # Create a Word document and save the extracted text
            doc = Document()
            doc.add_heading('Extracted Text', level=1)
            doc.add_paragraph(extracted_text.strip())
            
            # Save the document
            word_document_path = r"D:\Pictures\Screenshots\extracted_text.docx"
            doc.save(word_document_path)
            print(f"Word document saved as: {word_document_path}")

# Close the OpenCV window
cv2.destroyAllWindows()
